import sys
import os

import humanize
import docx

from wikinator.docxit import CommentBlock, build_numbering_cache, get_marker

# Not directly related to processing docs, but useful in development and understanding of the DOCX structures.

# def rando(length:int) -> str:
#    return ''.join(random.choice(string.ascii_lowercase) for i in range(length))

# class CommentBlock:
#     def __init__(self, comments: list[int]):
#         self.comments = comments
#         self.anchor = str(self.comments[0]) # just using first comment ID

#     def marker(self) -> str:
#         # generate a footnot marker for the comment-block
#         return "FIXME"

#     def generate_comment(self, doc:docx.Document) -> str:
#         """
#         Build the comment block in markdown,
#         as a footnote with formatting and time-ordered
#         comments
#         """
#         comment_str = "\n"
#         for comment_id in self.comments:
#             comment = doc.comments.get(comment_id)
#             if comment:
#                 comment_str += f"name: **{comment.author}**<br/>\n"
#                 comment_str += f"date: *{comment.timestamp}*<br/>\n"
#                 # note: the comment might contain styling. this strips that. might FIXME
#                 comment_str += "> " + comment.text
#                 comment_str += "\n---\n"

#         return comment_str

#     @staticmethod
#     def from_run(run:docx.text.run.Run):
#         """Returns a list of comments referenced in the run, or None"""
#         if not run:
#             return None

#         # scan the next elements (siblings) for commentRangeEnd
#         # and skipping empty runs.

#         comments = []
#         next = run._element.getnext()
#         while next is not None:
#             if isinstance(next, docx.oxml.text.run.CT_R):
#                 # skip empty text, break on non-empty
#                 if len(next.text) > 0:
#                     break
#             elif next.tag.endswith("commentRangeEnd"):
#                 # found a comment-end before the next text, store it
#                 comment_id = next.values()[0]
#                 comments.append(comment_id)
#             next = next.getnext()

#         if len(comments) > 0:
#             return CommentBlock(comments)
#         else:
#             return None


# def comments_in_run(run:docx.text.run.Run) -> CommentBlock | None:
#     """Returns a list of comments referenced in the run, or None"""
#     if not run:
#         return None

#     # scan the next elements (siblings) for commentRangeEnd
#     # and skipping empty runs.

#     comments = []
#     next = run._element.getnext()
#     while next is not None:
#         if isinstance(next, docx.oxml.text.run.CT_R):
#             # skip empty text, break on non-empty
#             if len(next.text) > 0:
#                 break
#         elif next.tag.endswith("commentRangeEnd"):
#             # found a comment-end before the next text, store it
#             comment_id = next.values()[0]
#             comments.append(comment_id)
#         next = next.getnext()

#     if len(comments) > 0:
#         return CommentBlock(comments)
#     else:
#         return None


def docx_dump(filename:str):
    if not filename.endswith(".docx"):
        print(f".... skipping {filename}")
        return

    # Dump a bunch of info about a docx file
    doc = docx.Document(filename)

    print(f"file: {filename}, {humanize.naturalsize(os.path.getsize(filename))}")

    numberingCache = build_numbering_cache(doc)


    props = doc.core_properties
    for name in dir(props):
        if not name.startswith("_"):
            value = getattr(props, name, None)
            if value:
                print(f"  > {name}: {getattr(props, name, "<not set>")}")

    # paragraphs
    for paragraph in doc.paragraphs:
        print("\n---\n")

        #marker = get_list_marker(doc, paragraph)
        #if marker:
        #    print(f"  MARKER: {marker}")

        numId, ilvl = get_marker(paragraph)
        if numId:
            numbering = numberingCache.get(numId, ilvl)
            print(f"NUMB: {numId}, {ilvl}: {numbering}")

            if numbering:
                print(f"NUMBERING: {numbering.format} - {numbering.style}")

        for run in paragraph.runs:
            if len(run.text) > 0:
                #print(f"{run.text}")
                comments = CommentBlock.from_run(run)
                if comments:
                    comment_str = comments.comments_from_doc(doc)
                    print(f"-- comments: {comment_str}")

    tables = doc.tables
    if len(tables) > 0:
        print(f"> {len(tables)} tables")

    # comments = doc.comments
    # if comments:
    #     for comment in comments:
    #         print(f"-- Comment #{comment.comment_id}, from: {comment.author}/{comment.initials}, ts: {comment.timestamp}")
    #         print(comment.text)
    #         #for p in comment.paragraphs:
    #         #    print(f"        --: {p.text}")

    shapes = doc.inline_shapes
    if shapes:
        print(f"> {len(shapes)} shapes")

    settings = doc.settings
    print(f"> SETTINGS {vars(settings)}")

    for style in doc.styles:
        print(f"  STYLE {style}")



    ### EXTRACT
    test_id = 4
    numbering = doc.part.numbering_part.numbering_definitions
    if numbering:
        for zorb in numbering._numbering:
            if zorb.tag.endswith("abstractNum"):
                abId = zorb.attrib.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}abstractNumId")
                if test_id == abId:
                    # numFmt: "decimal", "lowerLetter", "lowerRoman", "bullet"

                    print(f"abstractNum {zorb} - ID:{abId}")

            else:
                print(f"NUMBERING: {zorb}")






def main():
    for filename in sys.argv:
        docx_dump(filename)


if __name__ == "__main__":
    main()